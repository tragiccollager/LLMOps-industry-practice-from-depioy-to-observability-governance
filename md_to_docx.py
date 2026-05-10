"""
Markdown 转 Word 文档转换器
用于将实验报告 Markdown 文件转换为 Word 文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os


def set_chinese_font(run, font_name='SimSun', font_size=12, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def parse_markdown(md_content):
    """解析 Markdown 内容"""
    lines = md_content.split('\n')
    elements = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 标题
        if line.startswith('# '):
            elements.append(('h1', line[2:]))
        elif line.startswith('## '):
            elements.append(('h2', line[3:]))
        elif line.startswith('### '):
            elements.append(('h3', line[4:]))
        elif line.startswith('#### '):
            elements.append(('h4', line[5:]))
        
        # 表格
        elif line.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            elements.append(('table', table_lines))
            continue
        
        # 代码块
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append(('code', '\n'.join(code_lines)))
        
        # 列表项
        elif line.startswith('- ') or line.startswith('* '):
            elements.append(('bullet', line[2:]))
        elif re.match(r'^\d+\.\s', line):
            match = re.match(r'^(\d+)\.\s(.+)$', line)
            if match:
                elements.append(('number', match.group(2), int(match.group(1))))
        
        # 引用
        elif line.startswith('>'):
            elements.append(('quote', line[1:].strip()))
        
        # 分隔线
        elif line == '---' or line == '***':
            elements.append(('hr', ''))
        
        # 普通段落
        else:
            elements.append(('paragraph', line))
        
        i += 1
    
    return elements


def add_table(doc, table_lines):
    """添加表格"""
    if len(table_lines) < 3:
        return
    
    # 解析表头
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # 创建表格
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 填充表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, 'SimHei', 10, bold=True)
    
    # 填充数据行（跳过分隔行）
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells and any(cells):
            row_cells = table.add_row().cells
            for i, cell in enumerate(cells):
                if i < len(row_cells):
                    row_cells[i].text = cell
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            set_chinese_font(run, 'SimSun', 10)


def markdown_to_docx(md_file, docx_file):
    """将 Markdown 文件转换为 Word 文档"""
    
    # 读取 Markdown 内容
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 解析 Markdown
    elements = parse_markdown(md_content)
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)
    
    # 处理元素
    i = 0
    while i < len(elements):
        elem_type, *content = elements[i]
        
        if elem_type == 'h1':
            p = doc.add_heading(content[0], level=1)
            for run in p.runs:
                set_chinese_font(run, 'SimHei', 18, bold=True)
        
        elif elem_type == 'h2':
            p = doc.add_heading(content[0], level=2)
            for run in p.runs:
                set_chinese_font(run, 'SimHei', 16, bold=True)
        
        elif elem_type == 'h3':
            p = doc.add_heading(content[0], level=3)
            for run in p.runs:
                set_chinese_font(run, 'SimHei', 14, bold=True)
        
        elif elem_type == 'h4':
            p = doc.add_heading(content[0], level=4)
            for run in p.runs:
                set_chinese_font(run, 'SimHei', 12, bold=True)
        
        elif elem_type == 'table':
            add_table(doc, content[0])
            doc.add_paragraph()
        
        elif elem_type == 'code':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(content[0])
            set_chinese_font(run, 'Courier New', 9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        elif elem_type == 'bullet':
            p = doc.add_paragraph(content[0], style='List Bullet')
            for run in p.runs:
                set_chinese_font(run, 'SimSun', 12)
        
        elif elem_type == 'number':
            p = doc.add_paragraph(content[0], style='List Number')
            for run in p.runs:
                set_chinese_font(run, 'SimSun', 12)
        
        elif elem_type == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(content[0])
            set_chinese_font(run, 'SimSun', 11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        elif elem_type == 'hr':
            doc.add_paragraph('_' * 50)
        
        elif elem_type == 'paragraph':
            # 检查是否是粗体
            text = content[0]
            p = doc.add_paragraph()
            
            # 处理 **粗体**
            parts = re.split(r'\*\*(.+?)\*\*', text)
            for j, part in enumerate(parts):
                if j % 2 == 1:  # 奇数索引是粗体内容
                    run = p.add_run(part)
                    set_chinese_font(run, 'SimHei', 12, bold=True)
                else:
                    if part:
                        run = p.add_run(part)
                        set_chinese_font(run, 'SimSun', 12)
        
        i += 1
    
    # 保存文档
    doc.save(docx_file)
    print(f"[OK] Word 文档已生成: {docx_file}")


if __name__ == "__main__":
    md_file = "实验报告.md"
    docx_file = "2410311228_邬顺豪_LLMOps实验报告.docx"
    
    if os.path.exists(md_file):
        markdown_to_docx(md_file, docx_file)
    else:
        print(f"❌ 找不到文件: {md_file}")
